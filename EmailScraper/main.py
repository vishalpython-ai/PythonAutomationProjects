"""
Email, Phone, and Name Scraper - Single Row Format

This script extracts Name, Email, and Phone number from documents (PDF or Word)
where each contact is stored in a single row, separated by commas:

Example row:
John Doe, john.doe@example.com, +91-9876543210

It saves the extracted data in multiple formats:
1. TXT file
2. Excel (.xlsx)
3. Word (.docx)
4. PDF

Important Notes:
- Emails are validated with regex.
- Phone numbers can be Indian or international formats.
- Each line in the input file should have exactly three fields: Name, Email, Phone.
- Input files are read from 'input_files' folder.
- Output files are saved in 'output_files' folder.
"""

import os
from pathlib import Path
import re
import openpyxl
import docx
from reportlab.pdfgen import canvas
from PyPDF2 import PdfReader

# ------------------------------ Directories ------------------------------
BASE_DIR = Path(__file__).parent
INPUT_DIR = BASE_DIR / 'input_files'
OUTPUT_DIR = BASE_DIR / 'output_files'
OUTPUT_DIR.mkdir(exist_ok=True)

# ------------------------------ Regex Patterns ------------------------------
EMAIL_REGEX = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-z]{2,}")
PHONE_REGEX = re.compile(r"(\+?\d{1,4}[\s-]?)?(?:\(?\d{2,4}\)?[\s-]?)?\d{6,10}")

# ------------------------------ Functions ------------------------------
def extract_from_pdf(file_path):
    text = ""
    reader = PdfReader(str(file_path))
    for page in reader.pages:
        text += page.extract_text() + "\n"
    return text

def extract_from_word(file_path):
    doc = docx.Document(str(file_path))
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

def extract_contacts(text):
    contacts = []
    for line in text.splitlines():
        parts = [p.strip() for p in line.split(',')]
        if len(parts) != 3:
            continue
        name, email, phone = parts
        if EMAIL_REGEX.fullmatch(email) and PHONE_REGEX.fullmatch(phone.replace(" ", "")):
            contacts.append((name, email, phone))
    return contacts

def save_to_text(contacts, filename):
    with open(filename, "w") as f:
        f.write("Name | Email | Phone\n")
        f.write("-" * 40 + "\n")
        for name, email, phone in contacts:
            f.write(f"{name} | {email} | {phone}\n")

def save_to_excel(contacts, filename):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Contacts"
    ws.append(["Name", "Email", "Phone"])
    for name, email, phone in contacts:
        ws.append([name, email, phone])
    wb.save(filename)

def save_to_word(contacts, filename):
    doc = docx.Document()
    doc.add_heading("Extracted Contacts", level=1)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Email'
    hdr_cells[2].text = 'Phone'
    for name, email, phone in contacts:
        row_cells = table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = email
        row_cells[2].text = phone
    doc.save(filename)

def save_to_pdf(contacts, filename):
    c = canvas.Canvas(str(filename))
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, 800, "Extracted Contacts")
    c.setFont("Helvetica", 12)
    y = 770
    c.drawString(50, y, "Name | Email | Phone")
    y -= 20
    c.drawString(50, y, "-"*70)
    y -= 20
    for name, email, phone in contacts:
        line = f"{name} | {email} | {phone}"
        c.drawString(50, y, line)
        y -= 20
        if y < 50:
            c.showPage()
            y = 800
    c.save()

# ------------------------------ Main Code ------------------------------
if __name__ == "__main__":
    all_contacts = []

    for file in INPUT_DIR.iterdir():
        if file.suffix.lower() == ".pdf":
            text = extract_from_pdf(file)
        elif file.suffix.lower() in [".docx", ".doc"]:
            text = extract_from_word(file)
        else:
            continue
        contacts = extract_contacts(text)
        all_contacts.extend(contacts)

    # Remove duplicates
    all_contacts = list(set(all_contacts))

    # Save outputs
    save_to_text(all_contacts, OUTPUT_DIR / "contacts.txt")
    save_to_excel(all_contacts, OUTPUT_DIR / "contacts.xlsx")
    save_to_word(all_contacts, OUTPUT_DIR / "contacts.docx")
    save_to_pdf(all_contacts, OUTPUT_DIR / "contacts.pdf")

    print(f"Extraction complete. Files saved in: {OUTPUT_DIR}")
